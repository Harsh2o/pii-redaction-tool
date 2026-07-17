"""
redact.py

Main entry point for the PII Redaction Tool.

Usage:
    python redact.py --input "Red Herring Prospectus.docx" --output redacted_output.docx
    python redact.py --input doc.docx --output out.docx --dry-run
    python redact.py --input doc.docx --output out.docx --report summary_report.json --evaluate
"""

import argparse
import json
import logging
import os
import sys
import time
import tracemalloc
from datetime import datetime

import yaml
from docx import Document

from anonymizers import generate_replacement, get_replacement_cache
from recognizers import build_recognizer_registry
from tqdm import tqdm



def setup_logging(log_file: str):
    log_format = "%(asctime)s | %(levelname)-5s | %(message)s"
    date_fmt = "%Y-%m-%d %H:%M:%S"
    handlers = [
        logging.FileHandler(log_file, encoding="utf-8"),
        logging.StreamHandler(sys.stdout),
    ]
    logging.basicConfig(level=logging.INFO, format=log_format,
                        datefmt=date_fmt, handlers=handlers)
    
   
    logging.getLogger("presidio-analyzer").setLevel(logging.WARNING)
    logging.getLogger("presidio_analyzer").setLevel(logging.WARNING)
    
    return logging.getLogger(__name__)



def load_config(config_path: str) -> dict:
    defaults = {
        "redact": {
            "person_names": True, "emails": True, "phone_numbers": True,
            "company_names": True, "addresses": True, "cin_numbers": True,
            "pan_numbers": True, "aadhaar": True, "dates_of_birth": True,
            "ip_addresses": True, "urls": True, "credit_cards": True,
        },
        "settings": {
            "confidence_threshold": 0.45,
            "spacy_model": "en_core_web_lg",
            "faker_locale": "en_IN",
            "consistent_replacements": True,
            "log_file": "redaction.log",
        },
        "org_whitelist": [
            "SEBI", "NSE", "BSE", "RBI", "MCA", "NCLT", "NSDL", "CDSL",
            "IRDAI", "FEMA", "PMLA",
            "Securities and Exchange Board of India",
            "National Stock Exchange", "Bombay Stock Exchange",
            "Reserve Bank of India", "Ministry of Corporate Affairs",
        ]
    }
    if os.path.exists(config_path):
        with open(config_path, encoding="utf-8") as f:
            user_cfg = yaml.safe_load(f) or {}
      
        for section in defaults:
            if section in user_cfg:
                if isinstance(defaults[section], dict):
                    defaults[section].update(user_cfg[section])
                else:
                    defaults[section] = user_cfg[section]
    return defaults



def validate_inputs(input_path: str, output_path: str, dry_run: bool):
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"Input file not found: '{input_path}'")
    if not input_path.lower().endswith(".docx"):
        raise ValueError(f"Only .docx files are supported. Got: '{input_path}'")
    if os.path.getsize(input_path) == 0:
        raise ValueError(f"Input file is empty: '{input_path}'")
    try:
        Document(input_path)
    except Exception as e:
        raise RuntimeError(f"Could not open document (possibly corrupted): {e}")
    if not dry_run:
        out_dir = os.path.dirname(os.path.abspath(output_path)) or "."
        if not os.access(out_dir, os.W_OK):
            raise PermissionError(f"No write permission for: '{out_dir}'")



# Build the Presidio analyzer

def build_analyzer(config: dict, logger):
    from presidio_analyzer import AnalyzerEngine
    from presidio_analyzer.nlp_engine import NlpEngineProvider

    model_name = config["settings"]["spacy_model"]
    threshold  = config["settings"]["confidence_threshold"]

    logger.info(f"Loading spaCy model: {model_name}")

   
    ner_config = {"labels_to_ignore": [
        "FAC", "CARDINAL", "PRODUCT", "LAW", "WORK_OF_ART", 
        "MONEY", "QUANTITY", "ORDINAL", "PERCENT", "TIME", "EVENT", "LANGUAGE"
    ]}

  
    try:
        nlp_config = {
            "nlp_engine_name": "spacy",
            "models": [{"lang_code": "en", "model_name": model_name}],
            "ner_model_configuration": ner_config,
        }
        provider = NlpEngineProvider(nlp_configuration=nlp_config)
        nlp_engine = provider.create_engine()
    except Exception as e:
        logger.warning(f"Could not load {model_name}: {e}. Falling back to en_core_web_sm")
        nlp_config = {
            "nlp_engine_name": "spacy",
            "models": [{"lang_code": "en", "model_name": "en_core_web_sm"}],
            "ner_model_configuration": ner_config,
        }
        provider = NlpEngineProvider(nlp_configuration=nlp_config)
        nlp_engine = provider.create_engine()

    analyzer = AnalyzerEngine(
        nlp_engine=nlp_engine,
        default_score_threshold=threshold
    )

  
    for recognizer in build_recognizer_registry():
        analyzer.registry.add_recognizer(recognizer)

    logger.info(f"Analyzer ready | threshold={threshold}")
    return analyzer



# Detect entities in document to build document-derived deny list

def extract_document_entities(doc: Document, analyzer, logger) -> dict:
    """
    First pass: scan doc at high confidence to find names/orgs.
    Returns dict with 'persons' and 'orgs' sets.
    Used to supplement detection in the main pass.
    """
    persons = set()
    orgs    = set()

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        try:
            results = analyzer.analyze(
                text=text, language="en",
                score_threshold=0.75,
                entities=["PERSON", "ORGANIZATION"]
            )
            for r in results:
                entity_text = text[r.start:r.end].strip()
                if r.entity_type == "PERSON":
                    persons.add(entity_text)
                elif r.entity_type == "ORGANIZATION":
                    orgs.add(entity_text)
        except Exception:
            pass

        # Edge case: small spaCy model sometimes misses names with low confidence.
        # Hardcode extraction for known patterns like "Contact Person: Name"
        import re
        cp_matches = re.findall(r"Contact Person:\s*([A-Z][a-z]+(?:\s[A-Z][a-z]+)*)", text)
        for m in cp_matches:
            persons.add(m.strip())

    logger.info(f"Document-derived persons: {len(persons)}, orgs: {len(orgs)}")
    return {"persons": persons, "orgs": orgs}



# Core: redact a single paragraph (safe two-step strategy)

def _apply_replacements_to_paragraph(paragraph, replacements: list):
    """
    replacements: [(start, end, fake_text), ...]
    Two-step strategy:
      1. Try run-level replacement (preserves bold/italic per word)
      2. If ANY entity spans multiple runs, fallback: merge -> replace -> rebuild
    """
    if not replacements or not paragraph.runs:
        return

    =
    run_map = []
    cursor = 0
    for i, run in enumerate(paragraph.runs):
        run_map.append((cursor, cursor + len(run.text), i))
        cursor += len(run.text)

    
    all_single_run = True
    for (start, end, fake_text) in replacements:
        if not any(rs <= start and end <= re for (rs, re, _) in run_map):
            all_single_run = False
            break

    if all_single_run:
      
        for (start, end, fake_text) in sorted(replacements, reverse=True):
            for (rs, re, ri) in run_map:
                if rs <= start and end <= re:
                    run = paragraph.runs[ri]
                    ls  = start - rs
                    le  = end - rs
                    run.text = run.text[:ls] + fake_text + run.text[le:]
                    break
    else:
       
        full_text = paragraph.text
        for (start, end, fake_text) in sorted(replacements, reverse=True):
            full_text = full_text[:start] + fake_text + full_text[end:]

      
        if paragraph.runs:
            paragraph.runs[0].text = full_text
            for run in paragraph.runs[1:]:
                run.text = ""



# Analyze one paragraph and collect replacements

def _analyze_paragraph(text: str, analyzer, org_whitelist: set,
                        config: dict, entity_counts: dict) -> list:
    """Returns list of (start, end, fake_text) tuples."""
    if not text.strip():
        return []

    # Explicitly request entities  ensures ORGANIZATION is always checked
    requested_entities = [
        "PERSON", "EMAIL_ADDRESS", "PHONE_NUMBER", "LOCATION", "ORGANIZATION",
        "IP_ADDRESS", "CREDIT_CARD", "URL", "NRP",
        "IN_CIN", "IN_PAN", "IN_AADHAAR", "IN_GSTIN",
        "IN_PHONE_NUMBER", "IN_IFSC", "WEBSITE_URL", "DATE_OF_BIRTH",
    ]
    try:
        results = analyzer.analyze(text=text, language="en",
                                   entities=requested_entities)
    except Exception:
        return []

    if not results:
        return []

    replacements = []
    seen_spans   = set()
    redact_cfg   = config.get("redact", {})

    # Map Presidio entity types to config keys
    cfg_map = {
        "PERSON":           "person_names",
        "EMAIL_ADDRESS":    "emails",
        "PHONE_NUMBER":     "phone_numbers",
        "IN_PHONE_NUMBER":  "phone_numbers",
        "ORGANIZATION":     "company_names",
        "LOCATION":         "addresses",
        "GPE":              "addresses",
        "IP_ADDRESS":       "ip_addresses",
        "URL":              "urls",
        "WEBSITE_URL":      "urls",
        "CREDIT_CARD":      "credit_cards",
        "DATE_OF_BIRTH":    "dates_of_birth",
        "IN_CIN":           "cin_numbers",
        "IN_PAN":           "pan_numbers",
        "IN_AADHAAR":       "aadhaar",
        "IN_GSTIN":         "cin_numbers",
        "IN_IFSC":          "cin_numbers",
    }

    # Resolve overlaps (longest span wins for overlapping starts)
    results.sort(key=lambda x: (x.start, -(x.end - x.start)))
    resolved_results = []
    for r in results:
        if any(r.start < rr.end and rr.start < r.end for rr in resolved_results):
            continue
        resolved_results.append(r)

    for r in resolved_results:
        span = (r.start, r.end)

        entity_type = r.entity_type
        original    = text[r.start:r.end]

       
        cfg_key = cfg_map.get(entity_type)
        if cfg_key and not redact_cfg.get(cfg_key, True):
            continue

        if entity_type == "ORGANIZATION":
            if any(original.strip().lower() == w.lower() for w in org_whitelist):
                continue

     
        fake_text = generate_replacement(entity_type, original)
        replacements.append((r.start, r.end, fake_text))
        seen_spans.add(span)
        entity_counts[entity_type] = entity_counts.get(entity_type, 0) + 1

    return replacements





def _process_element_paragraphs(paragraphs, analyzer, org_whitelist,
                                  config, entity_counts):
    for para in paragraphs:
        text = para.text
        if not text.strip():
            continue
        reps = _analyze_paragraph(text, analyzer, org_whitelist,
                                   config, entity_counts)
        if reps:
            _apply_replacements_to_paragraph(para, reps)




def redact_document(input_path: str, output_path: str, config: dict,
                    logger, dry_run: bool = False) -> dict:

    logger.info(f"Loading document: {input_path}")
    doc = Document(input_path)

    para_count  = len(doc.paragraphs)
    table_count = len(doc.tables)
    logger.info(f"Document loaded: {para_count} paragraphs, {table_count} tables")

    analyzer     = build_analyzer(config, logger)
    org_whitelist = set(config.get("org_whitelist", []))

    
    logger.info("Running first pass to extract document-derived entities...")
    doc_entities = extract_document_entities(doc, analyzer, logger)

   
    if doc_entities:
        from presidio_analyzer import PatternRecognizer
        # Add persons
        if doc_entities.get("persons"):
            persons_list = list(doc_entities["persons"])
            person_rec = PatternRecognizer(
                supported_entity="PERSON",
                deny_list=persons_list,
                name="DocumentDerivedPersonRecognizer"
            )
            analyzer.registry.add_recognizer(person_rec)
            logger.info(f"Added {len(persons_list)} document-derived person names to recognizer")
        # Add orgs
        if doc_entities.get("orgs"):
            orgs_list = [o for o in doc_entities["orgs"]
                         if not any(o.strip().lower() == w.lower() for w in org_whitelist)]
            if orgs_list:
                org_rec = PatternRecognizer(
                    supported_entity="ORGANIZATION",
                    deny_list=orgs_list,
                    name="DocumentDerivedOrgRecognizer"
                )
                analyzer.registry.add_recognizer(org_rec)
                logger.info(f"Added {len(orgs_list)} document-derived org names to recognizer")

    entity_counts = {}

    if dry_run:
     
        logger.info("[DRY RUN] Analyzing document — no file will be written")
        print("\n" + "=" * 60)
        print("  [DRY RUN] Detected PII — no file written.")
        print("=" * 60)

        for para in doc.paragraphs:
            text = para.text
            if not text.strip():
                continue
            reps = _analyze_paragraph(text, analyzer, org_whitelist,
                                       config, entity_counts)
            for (start, end, fake_text) in reps:
                original = text[start:end]
                print(f"  {original:<40} ->  {fake_text}")

        total = sum(entity_counts.values())
        print("=" * 60)
        for etype, count in sorted(entity_counts.items()):
            print(f"  {etype:<25} {count}")
        print(f"\n  Total PII instances detected: {total}")
        print(f"  Run without --dry-run to generate the redacted document.")
        return entity_counts

   
    logger.info("Starting redaction...")

    
    logger.info("Processing body paragraphs...")
    for para in tqdm(doc.paragraphs, desc="Body Paragraphs", disable=dry_run):
        if para.text.strip():
            reps = _analyze_paragraph(para.text, analyzer, org_whitelist, config, entity_counts)
            if reps:
                _apply_replacements_to_paragraph(para, reps)

   
    logger.info("Processing tables...")
    for table in tqdm(doc.tables, desc="Tables", disable=dry_run):
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        reps = _analyze_paragraph(para.text, analyzer, org_whitelist, config, entity_counts)
                        if reps:
                            _apply_replacements_to_paragraph(para, reps)

    # Headers and footers
    for section in doc.sections:
        try:
            _process_element_paragraphs(section.header.paragraphs, analyzer,
                                         org_whitelist, config, entity_counts)
            for tbl in section.header.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        _process_element_paragraphs(cell.paragraphs, analyzer,
                                                     org_whitelist, config, entity_counts)
        except Exception:
            pass
        try:
            _process_element_paragraphs(section.footer.paragraphs, analyzer,
                                         org_whitelist, config, entity_counts)
        except Exception:
            pass

    # Log per-type counts
    total = sum(entity_counts.values())
    for etype, count in sorted(entity_counts.items()):
        logger.info(f"  Detected {count:>4} {etype}")
    logger.info(f"Total PII instances detected: {total}")

    # Save
    logger.info(f"Saving redacted document: {output_path}")
    doc.save(output_path)
    logger.info(f"Saved successfully.")

    return entity_counts



# Save JSON summary report

def save_summary_report(entity_counts: dict, input_path: str,
                         elapsed: float, peak_mb: float,
                         report_path: str, logger):
    report = {
        "document":                 os.path.basename(input_path),
        "timestamp":                datetime.now().isoformat(timespec="seconds"),
        "processing_time_seconds":  round(elapsed, 2),
        "memory_peak_mb":           round(peak_mb, 1),
        "total_pii_detected":       sum(entity_counts.values()),
        "by_type":                  dict(sorted(entity_counts.items())),
    }
    with open(report_path, "w", encoding="utf-8") as f:
        json.dump(report, f, indent=2)
    logger.info(f"Summary report saved -> {report_path}")
    return report


def parse_args():
    parser = argparse.ArgumentParser(
        description="PII Redaction Tool — redacts personally identifiable information from .docx files",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python redact.py --input "Red Herring Prospectus.docx" --output redacted_output.docx
  python redact.py --input doc.docx --output out.docx --dry-run
  python redact.py --input doc.docx --output out.docx --report summary.json --evaluate
        """
    )
    parser.add_argument("--input",     required=True,            help="Path to input .docx file")
    parser.add_argument("--output",    default="redacted_output.docx", help="Path to save redacted .docx")
    parser.add_argument("--config",    default="config.yaml",    help="Config file path (default: config.yaml)")
    parser.add_argument("--threshold", type=float, default=None, help="Override confidence threshold (0.0-1.0)")
    parser.add_argument("--high-accuracy", action="store_true",      help="Use en_core_web_lg model for better accuracy (uses more memory)")
    parser.add_argument("--report",    default=None,             help="Save JSON summary report to this path")
    parser.add_argument("--log",       default="redaction.log",  help="Log file path (default: redaction.log)")
    parser.add_argument("--evaluate",  action="store_true",      help="Run evaluator after redaction")
    parser.add_argument("--dry-run",   action="store_true",      help="Preview detections without writing output")
    parser.add_argument("--stats",     action="store_true",      help="Print execution stats at the end")
    return parser.parse_args()


def main():
    args = parse_args()
    logger = setup_logging(args.log)

    logger.info("=" * 55)
    logger.info("  PII Redaction Tool — starting")
    logger.info("=" * 55)

    # Load config
    config = load_config(args.config)
    if args.threshold is not None:
        config["settings"]["confidence_threshold"] = args.threshold
        logger.info(f"Threshold overridden via CLI: {args.threshold}")
    
    if args.high_accuracy:
        config["settings"]["spacy_model"] = "en_core_web_lg"
        logger.info("High accuracy mode enabled: using en_core_web_lg")

    # Validate inputs
    try:
        validate_inputs(args.input, args.output, args.dry_run)
    except (FileNotFoundError, ValueError, RuntimeError, PermissionError) as e:
        logger.error(str(e))
        sys.exit(1)

    # Track performance
    tracemalloc.start()
    start_time = time.time()

    try:
        entity_counts = redact_document(
            input_path  = args.input,
            output_path = args.output,
            config      = config,
            logger      = logger,
            dry_run     = args.dry_run,
        )
    except Exception as e:
        logger.error(f"Redaction failed: {e}")
        sys.exit(1)

    elapsed = time.time() - start_time
    _, peak = tracemalloc.get_traced_memory()
    peak_mb = peak / 1024 / 1024
    tracemalloc.stop()

    logger.info(f"Processing time: {elapsed:.2f}s | Peak memory: {peak_mb:.1f} MB")
    
    if args.stats:
        print("\n" + "=" * 55)
        print("  EXECUTION STATS")
        print("=" * 55)
        print(f"  Time taken: {elapsed:.2f} seconds")
        print(f"  Memory peak: {peak_mb:.1f} MB")
        print(f"  Total Entities Replaced: {sum(entity_counts.values())}")
        print("=" * 55)

    # Save summary report if requested
    if args.report and not args.dry_run:
        save_summary_report(entity_counts, args.input, elapsed, peak_mb,
                            args.report, logger)

    # Run evaluator if requested
    if args.evaluate and not args.dry_run:
        logger.info("Running evaluator...")
        import evaluator
        results = evaluator.evaluate(args.input, args.output)
        evaluator.print_results(results)
        evaluator.save_markdown_report(results)

    logger.info("Done.")


if __name__ == "__main__":
    main()
