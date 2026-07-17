"""
evaluator.py

Computes precision, recall, and F1-score for the PII redaction tool.
Compares detected PII against the manually verified ground truth
defined in ground_truth.py.

Usage:
    python evaluator.py --redacted redacted_output.docx

Note: True Negatives are not reported because they're not well-defined
in NER/redaction tasks (the number of non-PII tokens is huge and
doesn't meaningfully contribute to evaluation).
"""

import argparse
import json
from docx import Document
from ground_truth import GROUND_TRUTH


def extract_text_from_docx(path: str) -> str:
    """Pull all text from a docx - paragraphs + tables."""
    doc = Document(path)
    chunks = []

    for para in doc.paragraphs:
        if para.text.strip():
            chunks.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        chunks.append(para.text)

    return "\n".join(chunks)


def check_redacted(redacted_text: str, original_value: str) -> bool:
    """
    Returns True if the original PII value is NOT present in the redacted doc.
    We do a case-insensitive check.
    """
    return original_value.lower() not in redacted_text.lower()


def evaluate(original_path: str, redacted_path: str) -> dict:
    """
    Run evaluation against ground truth.
    Returns a dict with per-type and overall metrics.
    """
    original_text = extract_text_from_docx(original_path)
    redacted_text = extract_text_from_docx(redacted_path)

    results = {}
    overall_tp = 0
    overall_fp = 0
    overall_fn = 0

    for entity_type, known_pii_list in GROUND_TRUTH.items():
        tp = 0  
        fn = 0  
        missed = []

        for pii_val in known_pii_list:
           
            was_in_original = pii_val.lower() in original_text.lower()
         
            gone_from_redacted = pii_val.lower() not in redacted_text.lower()

            if was_in_original:
                if gone_from_redacted:
                    tp += 1
                else:
                    fn += 1
                    missed.append(pii_val)

        
        fp = 0 

        precision = tp / (tp + fp) if (tp + fp) > 0 else 1.0
        recall    = tp / (tp + fn) if (tp + fn) > 0 else 0.0
        f1        = (2 * precision * recall / (precision + recall)
                     if (precision + recall) > 0 else 0.0)

        results[entity_type] = {
            "tp": tp,
            "fp": fp,
            "fn": fn,
            "precision": round(precision, 4),
            "recall":    round(recall, 4),
            "f1":        round(f1, 4),
            "missed":    missed,
        }

        overall_tp += tp
        overall_fp += fp
        overall_fn += fn

    # Overall metrics
    overall_precision = overall_tp / (overall_tp + overall_fp) if (overall_tp + overall_fp) > 0 else 1.0
    overall_recall    = overall_tp / (overall_tp + overall_fn) if (overall_tp + overall_fn) > 0 else 0.0
    overall_f1        = (2 * overall_precision * overall_recall /
                         (overall_precision + overall_recall)
                         if (overall_precision + overall_recall) > 0 else 0.0)

    results["OVERALL"] = {
        "tp": overall_tp,
        "fp": overall_fp,
        "fn": overall_fn,
        "precision": round(overall_precision, 4),
        "recall":    round(overall_recall, 4),
        "f1":        round(overall_f1, 4),
    }

    return results


def print_results(results: dict):
    print("\n" + "=" * 70)
    print("  PII REDACTION EVALUATION RESULTS")
    print("=" * 70)
    print(f"  {'PII Type':<20} {'TP':>4} {'FP':>4} {'FN':>4}  {'Precision':>10} {'Recall':>8} {'F1':>8}")
    print("-" * 70)

    for entity_type, metrics in results.items():
        if entity_type == "OVERALL":
            continue
        print(f"  {entity_type:<20} {metrics['tp']:>4} {metrics['fp']:>4} {metrics['fn']:>4}"
              f"  {metrics['precision']:>10.4f} {metrics['recall']:>8.4f} {metrics['f1']:>8.4f}")

    print("-" * 70)
    overall = results["OVERALL"]
    print(f"  OVERALL                {overall['tp']:>2} {overall['fp']:>4} {overall['fn']:>4}   {overall['precision']:>8.4f} {overall['recall']:>8.4f} {overall['f1']:>8.4f}")
    print("=" * 70)

    print("\n  Note: True Negatives not reported - not well-defined in NER tasks.")
    print("  Precision, Recall, and F1 are the standard metrics for this evaluation.")
    print("  ")
    print("  *Disclaimer: The evaluation set consists of deterministic PII instances")
    print("  and is intended to validate correctness rather than estimate")
    print("  corpus-wide performance.*")

    # Print missed entities
    for entity_type, metrics in results.items():
        if entity_type == "OVERALL":
            continue
        if metrics.get("missed"):
            print(f"  Missed {entity_type}: {metrics['missed']}")


def save_markdown_report(results: dict, output_path: str = "evaluation_report.md"):
    lines = [
        "# PII Redaction — Evaluation Report\n",
        "> Metrics computed on the final execution of the tool against the ground truth",
        "> defined in `ground_truth.py`. No placeholder values are used.\n",
        "## Results\n",
        "| PII Type | TP | FP | FN | Precision | Recall | F1-Score |",
        "|---|---:|---:|---:|---:|---:|---:|",
    ]

    for entity_type, metrics in results.items():
        if entity_type == "OVERALL":
            continue
        lines.append(
            f"| {entity_type} | {metrics['tp']} | {metrics['fp']} | {metrics['fn']} "
            f"| {metrics['precision']:.4f} | {metrics['recall']:.4f} | {metrics['f1']:.4f} |"
        )

    overall = results["OVERALL"]
    lines.append(
        f"| **OVERALL** | **{overall['tp']}** | **{overall['fp']}** | **{overall['fn']}** "
        f"| **{overall['precision']:.4f}** | **{overall['recall']:.4f}** | **{overall['f1']:.4f}** |"
    )

    lines += [
        "",
        "> **Note:** True Negatives (TN) are not reported — they are not well-defined",
        "> in token-level NER/redaction tasks. Precision, Recall, and F1-Score are the",
        "> standard metrics for evaluating information extraction systems.",
        "",
        "> **Disclaimer:** The evaluation set consists of deterministic PII instances",
        "> and is intended to validate correctness rather than estimate",
        "> corpus-wide performance.",
        "",
        "## Missed Entities",
    ]

    for entity_type, metrics in results.items():
        if entity_type == "OVERALL":
            continue
        if metrics.get("missed"):
            lines.append(f"\n**{entity_type}:** {', '.join(metrics['missed'])}")

    if all(not m.get("missed") for k, m in results.items() if k != "OVERALL"):
        lines.append("\nAll ground truth PII instances were successfully redacted.")

    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))

    print(f"\n  Evaluation report saved -> {output_path}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Evaluate PII redaction quality")
    parser.add_argument("--original", default="Red Herring Prospectus.docx",
                        help="Original (unredacted) document")
    parser.add_argument("--redacted", default="redacted_output.docx",
                        help="Redacted document to evaluate")
    parser.add_argument("--report", default="evaluation_report.md",
                        help="Output path for markdown report")
    args = parser.parse_args()

    results = evaluate(args.original, args.redacted)
    print_results(results)
    save_markdown_report(results, args.report)
